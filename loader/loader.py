"""
Universal Document Loader
Handles: PDF (text), PDF (scanned), DOCX, Images, Email Attachments, TXT
"""

import os
import pytesseract

# ============ WINDOWS TESSERACT PATH ============

pytesseract.pytesseract.pytesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
# ================================================

import PyPDF2
from docx import Document
from PIL import Image
import pdf2image
import base64
from typing import Dict, List
import email
from email import policy
from io import BytesIO

class DocumentLoader:
    """Unified loader for all document types."""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.png', '.jpg', '.jpeg', '.bmp', '.eml', '.txt']
    
    # ============ PDF HANDLING ============
    
    def _is_scanned_pdf(self, file_path: str) -> bool:
        """Check if PDF is scanned (image-based)."""
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                if len(reader.pages) == 0:
                    return True
                text = reader.pages[0].extract_text()
                result = len(text.strip()) < 100  # More lenient threshold
                print(f"   PDF text check: {len(text)} chars → {'Scanned' if result else 'Text-based'}")
                return result
        except:
            return True
    
    def _extract_text_pdf(self, file_path: str) -> str:
        """Extract text from text-based PDF using PyPDF2."""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            print(f"❌ PDF text extraction failed: {e}")
        return text
    
    def _extract_scanned_pdf_ocr(self, file_path: str) -> str:
        """
        Extract text from scanned PDF using OCR (Tesseract).
        Converts PDF pages to images → OCR → text
        """
        try:
            images = pdf2image.convert_from_path(file_path)
            all_text = []
            
            for i, image in enumerate(images):
                print(f"🔍 OCR processing page {i+1}/{len(images)}...")
                text = pytesseract.image_to_string(image)
                all_text.append(text)
            
            return "\n\n".join(all_text)
        except Exception as e:
            print(f"❌ Scanned PDF OCR failed: {e}")
            return ""
    
    def load_pdf(self, file_path: str) -> str:
        """
        Load PDF with automatic OCR detection.
        - Text PDF → PyPDF2
        - Scanned PDF → OCR (Tesseract)
        """
        print(f"📄 Loading PDF: {file_path}")
        
        # Try normal extraction first
        text = self._extract_text_pdf(file_path)
        
        # If very little text, use OCR
        if len(text.strip()) < 100:
            print("⚠️ PDF appears scanned, using OCR...")
            text = self._extract_scanned_pdf_ocr(file_path)
        
        return text if text else ""
    
    # ============ DOCX HANDLING ============
    
    def load_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        print(f"📄 Loading DOCX: {file_path}")
        text = ""
        try:
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
        except Exception as e:
            print(f"❌ DOCX extraction failed: {e}")
        
        return text
    
    # ============ IMAGE HANDLING ============
    
    def load_image(self, file_path: str) -> str:
        """
        Extract text from image using OCR.
        Supports: PNG, JPG, JPEG, BMP, TIFF, etc.
        
        IMPROVED: Better OCR config, error handling, and debugging
        """
        print(f"🖼️ Loading Image: {file_path}")
        try:
            # Open image
            image = Image.open(file_path)
            print(f"   ✅ Image opened: {image.size}, Mode: {image.mode}")
            
            # Convert to RGB if needed (important for OCR)
            if image.mode != 'RGB':
                print(f"   Converting {image.mode} → RGB")
                image = image.convert('RGB')
            
            # Save converted image temporarily for better OCR
            temp_rgb_path = file_path.replace('.png', '_rgb.png').replace('.jpg', '_rgb.jpg')
            image.save(temp_rgb_path)
            print(f"   Saved RGB version for OCR")
            
            # Extract text with optimized config
            print(f"   🔍 Running Tesseract OCR (Config: --psm 6)...")
            text = pytesseract.image_to_string(
                image,
                config='--psm 6'  # Assume single uniform block of text
            )
            
            print(f"   ✅ OCR Complete: {len(text)} chars extracted")
            
            # If very little text, try with different PSM
            if len(text.strip()) < 50:
                print(f"   ⚠️ Low text detected ({len(text)} chars), trying PSM 3...")
                text = pytesseract.image_to_string(
                    image,
                    config='--psm 3'  # Auto page segmentation
                )
                print(f"   ✅ Retry OCR: {len(text)} chars")
            
            # Clean up temp file
            try:
                if os.path.exists(temp_rgb_path):
                    os.remove(temp_rgb_path)
            except:
                pass
            
            result = text.strip() if text else ""
            
            if result:
                print(f"   ✅ Success! Extracted text preview: {result[:100]}...")
            else:
                print(f"   ⚠️ Warning: No text extracted from image")
            
            return result
        
        except FileNotFoundError:
            print(f"❌ Image file not found: {file_path}")
            return ""
        
        except Exception as e:
            print(f"❌ Image OCR failed: {e}")
            import traceback
            traceback.print_exc()
            return ""
    
    # ============ EMAIL HANDLING ============
    
    def load_email(self, file_path: str) -> Dict[str, any]:
        """
        Extract text from .eml (email) file.
        Returns: dict with subject, body, attachments
        """
        print(f"📧 Loading Email: {file_path}")
        
        try:
            with open(file_path, 'rb') as f:
                msg = email.message_from_binary_file(f, policy=policy.default)
            
            email_data = {
                "subject": msg.get('Subject', ''),
                "from": msg.get('From', ''),
                "to": msg.get('To', ''),
                "body": "",
                "attachments": []
            }
            
            # Extract body
            if msg.is_multipart():
                for part in msg.iter_parts():
                    if part.get_content_type() == "text/plain":
                        email_data["body"] += part.get_content()
                    elif part.get_content_type() == "text/html":
                        email_data["body"] += part.get_content()
                    elif part.get_content_disposition() == "attachment":
                        email_data["attachments"].append(part.get_filename())
            else:
                email_data["body"] = msg.get_content()
            
            print(f"   ✅ Email parsed: Subject='{email_data['subject'][:50]}...', Body={len(email_data['body'])} chars")
            
            return email_data
        
        except Exception as e:
            print(f"❌ Email parsing failed: {e}")
            return {"error": str(e)}
    
    # ============ TXT HANDLING ============
    
    def load_txt(self, file_path: str) -> str:
        """Extract text from .txt file."""
        print(f"📝 Loading TXT: {file_path}")
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            print(f"   ✅ TXT loaded: {len(text)} chars")
            return text if text else ""
        except Exception as e:
            print(f"❌ TXT loading failed: {e}")
            return ""
    
    # ============ UNIVERSAL LOADER ============
    
    def load_document(self, file_path: str) -> str:
        """
        Auto-detect format and load document.
        
        Supported formats:
        - PDF (text & scanned)
        - DOCX
        - Images (PNG, JPG, BMP)
        - Email (.eml)
        - Text (.txt)
        """
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.pdf':
            return self.load_pdf(file_path)
        elif ext == '.docx':
            return self.load_docx(file_path)
        elif ext in ['.png', '.jpg', '.jpeg', '.bmp', '.tiff']:
            return self.load_image(file_path)
        elif ext == '.eml':
            email_data = self.load_email(file_path)
            return f"{email_data.get('subject', '')}\n{email_data.get('body', '')}"
        elif ext == '.txt':
            return self.load_txt(file_path)
        else:
            raise ValueError(f"❌ Unsupported format: {ext}. Supported: {self.supported_formats}")


# ============ BACKWARD COMPATIBILITY FUNCTIONS ============

def load_document(file_path: str) -> str:
    """
    Auto-detect format and load document.
    (Backward compatible wrapper)
    """
    loader = DocumentLoader()
    return loader.load_document(file_path)


def load_pdf(file_path: str) -> str:
    """Load PDF file."""
    loader = DocumentLoader()
    return loader.load_pdf(file_path)


def load_docx(file_path: str) -> str:
    """Load DOCX file."""
    loader = DocumentLoader()
    return loader.load_docx(file_path)


def load_image(file_path: str) -> str:
    """Load image file."""
    loader = DocumentLoader()
    return loader.load_image(file_path)


def load_email(file_path: str) -> Dict:
    """Load email file."""
    loader = DocumentLoader()
    return loader.load_email(file_path)


def load_txt(file_path: str) -> str:
    """Load text file."""
    loader = DocumentLoader()
    return loader.load_txt(file_path)